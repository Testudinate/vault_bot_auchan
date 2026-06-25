#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════╗
║  Яндекс Диск → Obsidian Vault Sync                           ║
║                                                              ║
║  • Подключается через Яндекс Диск REST API (OAuth)           ║
║  • Скачивает все файлы рекурсивно                            ║
║  • Читает txt, md, json, csv, docx, pdf, xlsx               ║
║  • Инкрементальное обновление — только новые/изменённые      ║
║  • Запускается вручную или из Telegram бота                  ║
╚══════════════════════════════════════════════════════════════╝

КАК ПОЛУЧИТЬ OAUTH ТОКЕН:
  1. Зайдите на https://oauth.yandex.ru/
  2. Создайте приложение → выберите права "Яндекс Диск"
  3. Или используйте готовый токен из настроек Яндекс Диска

УСТАНОВКА:
  pip install requests python-docx PyMuPDF openpyxl

ЗАПУСК:
  python disk_sync.py                    # полная синхронизация
  python disk_sync.py --folder /Документы  # только одна папка
  python disk_sync.py --status           # показать статистику
"""

# ══════════════════════════════════════════════
#  НАСТРОЙКИ
# ══════════════════════════════════════════════
YANDEX_DISK_TOKEN = "y0__wgBEOa1kqaq94ACGNrsQCDwpemaFzIqPqVwryOFEZ5YBOr7p-3eS8J7"  # ← oauth.yandex.ru
VAULT_PATH        = "/Users/ods/Downloads/AI/ObsidianVault"
DISK_ROOT_FOLDER  = "/"        # "/" = весь диск, или "/Работа" для конкретной папки
MAX_FILE_SIZE_MB  = 10         # пропускать файлы больше этого размера
# ══════════════════════════════════════════════

import os, re, json, argparse, logging
from pathlib import Path
from datetime import datetime
from typing import Optional

import requests

logging.basicConfig(
    format="%(asctime)s [%(levelname)s] %(message)s",
    level=logging.INFO
)
log = logging.getLogger(__name__)

DISK_API     = "https://cloud-api.yandex.net/v1/disk"
STATE_FILE   = ".disk_sync_state.json"
VAULT_SECTION = "03_YandexDisk"

# Поддерживаемые форматы
TEXT_FORMATS = {".txt", ".md", ".json", ".csv", ".html", ".xml", ".yaml", ".yml"}
DOC_FORMATS  = {".docx", ".pdf", ".xlsx", ".xls", ".pptx"}
ALL_FORMATS  = TEXT_FORMATS | DOC_FORMATS


# ──────────────────────────────────────────────
#  ЯНДЕКС ДИСК API
# ──────────────────────────────────────────────

class YandexDiskAPI:
    def __init__(self, token: str):
        self.token   = token
        self.headers = {"Authorization": f"OAuth {token}"}
        self._check_token()

    def _check_token(self):
        resp = requests.get(f"{DISK_API}/", headers=self.headers, timeout=10)
        if resp.status_code == 401:
            raise ValueError("❌ Неверный токен Яндекс Диска. Проверьте YANDEX_DISK_TOKEN.")
        resp.raise_for_status()
        info = resp.json()
        log.info(f"✅ Подключён к Яндекс Диску: {info.get('user', {}).get('login', '?')}")
        log.info(f"   Занято: {info.get('used_space', 0)//1024//1024} МБ "
                 f"/ {info.get('total_space', 0)//1024//1024} МБ")

    def list_folder(self, path: str, limit: int = 100) -> list:
        """Возвращает список файлов и папок по пути."""
        items  = []
        offset = 0
        while True:
            resp = requests.get(
                f"{DISK_API}/resources",
                headers=self.headers,
                params={"path": path, "limit": limit, "offset": offset,
                        "fields": "name,path,type,size,modified,md5,_embedded"},
                timeout=30
            )
            if resp.status_code == 404:
                log.warning(f"Папка не найдена: {path}")
                return []
            resp.raise_for_status()
            data     = resp.json()
            embedded = data.get("_embedded", {})
            batch    = embedded.get("items", [])
            items.extend(batch)
            if len(batch) < limit:
                break
            offset += limit
        return items

    def list_all_files(self, root: str = "/") -> list:
        """Рекурсивно обходит все папки и возвращает все файлы."""
        all_files = []
        queue     = [root]
        visited   = set()

        while queue:
            folder = queue.pop(0)
            if folder in visited:
                continue
            visited.add(folder)

            log.info(f"  📂 Сканирую: {folder}")
            items = self.list_folder(folder)

            for item in items:
                if item["type"] == "dir":
                    queue.append(item["path"])
                elif item["type"] == "file":
                    ext = Path(item["name"]).suffix.lower()
                    if ext in ALL_FORMATS:
                        size_mb = item.get("size", 0) / 1024 / 1024
                        if size_mb <= MAX_FILE_SIZE_MB:
                            all_files.append(item)
                        else:
                            log.info(f"  ⏭  Пропущен (>{MAX_FILE_SIZE_MB}МБ): {item['name']}")

        log.info(f"  Найдено файлов: {len(all_files)}")
        return all_files

    def download_file(self, disk_path: str) -> bytes:
        """Скачивает файл и возвращает байты."""
        # Получаем ссылку для скачивания
        resp = requests.get(
            f"{DISK_API}/resources/download",
            headers=self.headers,
            params={"path": disk_path},
            timeout=30
        )
        resp.raise_for_status()
        url = resp.json()["href"]

        # Скачиваем файл
        file_resp = requests.get(url, timeout=60)
        file_resp.raise_for_status()
        return file_resp.content

    def get_disk_info(self) -> dict:
        resp = requests.get(f"{DISK_API}/", headers=self.headers, timeout=10)
        resp.raise_for_status()
        return resp.json()


# ──────────────────────────────────────────────
#  КОНВЕРТЕРЫ ФОРМАТОВ
# ──────────────────────────────────────────────

def extract_text_from_docx(data: bytes) -> str:
    try:
        import io
        from docx import Document
        doc   = Document(io.BytesIO(data))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                paras.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(paras)
    except ImportError:
        return "[docx: установите python-docx]"
    except Exception as e:
        return f"[docx ошибка: {e}]"


def extract_text_from_pdf(data: bytes) -> str:
    try:
        import io, fitz  # PyMuPDF
        doc   = fitz.open(stream=data, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except ImportError:
        return "[pdf: установите PyMuPDF: pip install PyMuPDF]"
    except Exception as e:
        return f"[pdf ошибка: {e}]"


def extract_text_from_xlsx(data: bytes) -> str:
    try:
        import io, openpyxl
        wb    = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"## Лист: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if row_text.strip(" |"):
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        return "[xlsx: установите openpyxl: pip install openpyxl]"
    except Exception as e:
        return f"[xlsx ошибка: {e}]"


def extract_text(filename: str, data: bytes) -> str:
    """Извлекает текст из файла любого поддерживаемого формата."""
    if not data:
        return "[Файл пуст]"
    ext = Path(filename).suffix.lower()

    if ext in TEXT_FORMATS:
        for enc in ["utf-8", "cp1251", "latin-1"]:
            try:
                return data.decode(enc)
            except Exception:
                continue
        return data.decode("utf-8", errors="replace")

    elif ext == ".docx":
        return extract_text_from_docx(data)
    elif ext == ".pdf":
        return extract_text_from_pdf(data)
    elif ext in {".xlsx", ".xls"}:
        return extract_text_from_xlsx(data)
    elif ext == ".pptx":
        return "[pptx: конвертация не поддерживается]"
    else:
        return data.decode("utf-8", errors="replace")


# ──────────────────────────────────────────────
#  СОСТОЯНИЕ (инкрементальная синхронизация)
# ──────────────────────────────────────────────

def load_sync_state() -> dict:
    try:
        return json.loads(Path(STATE_FILE).read_text(encoding="utf-8"))
    except Exception:
        return {"files": {}, "last_sync": None, "total_synced": 0}

def save_sync_state(state: dict):
    Path(STATE_FILE).write_text(
        json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

def file_changed(item: dict, state: dict) -> bool:
    """Проверяет изменился ли файл с последней синхронизации."""
    path     = item["path"]
    modified = item.get("modified", "")
    md5      = item.get("md5", "")
    cached   = state["files"].get(path, {})

    if not cached:
        return True  # новый файл
    if cached.get("md5") and md5 and cached["md5"] != md5:
        return True  # изменился
    if cached.get("modified") != modified:
        return True  # дата изменения другая
    return False


# ──────────────────────────────────────────────
#  СОХРАНЕНИЕ В VAULT
# ──────────────────────────────────────────────

def safe_path(disk_path: str) -> Path:
    """Конвертирует путь на диске в безопасный путь для файловой системы."""
    # Убираем ведущий слэш
    clean = disk_path.lstrip("/")
    # Заменяем небезопасные символы
    clean = re.sub(r'[<>:"|?*]', '_', clean)
    return Path(clean)


def save_to_vault(vault: Path, item: dict, text: str) -> Path:
    """Сохраняет файл с Диска как Markdown заметку в vault."""
    disk_path = item["path"]
    filename  = item["name"]
    modified  = item.get("modified", "")
    size      = item.get("size", 0)
    ext       = Path(filename).suffix.lower()

    # Структура папок повторяет Диск
    rel_path  = safe_path(disk_path)
    out_name  = rel_path.with_suffix(".md")
    out_path  = vault / VAULT_SECTION / out_name
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Форматируем дату
    try:
        dt       = datetime.fromisoformat(modified.replace("Z", "+00:00"))
        date_str = dt.strftime("%Y-%m-%d")
        date_fmt = dt.strftime("%d.%m.%Y %H:%M")
    except Exception:
        date_str = "unknown"
        date_fmt = modified

    # Emoji по типу файла
    type_emoji = {
        ".txt":"📄", ".md":"📝", ".json":"🔧", ".csv":"📊",
        ".docx":"📋", ".pdf":"📑", ".xlsx":"📊", ".html":"🌐",
    }.get(ext, "📄")

    md = f"""---
type: disk_file
disk_path: "{disk_path}"
original_name: "{filename}"
file_type: "{ext}"
date: {date_str}
modified: "{modified}"
size_bytes: {size}
source: yandex_disk_api
---

# {type_emoji} {filename}

**Путь на диске:** `{disk_path}`
**Изменён:** {date_fmt}
**Размер:** {size // 1024} КБ

---

{text[:50000]}
"""
    out_path.write_text(md, encoding="utf-8")
    return out_path


# ──────────────────────────────────────────────
#  ГЛАВНАЯ ФУНКЦИЯ СИНХРОНИЗАЦИИ
# ──────────────────────────────────────────────

def sync_disk(root_folder: str = "/", force: bool = False,
              progress_callback=None) -> dict:
    """
    Синхронизирует Яндекс Диск с Obsidian vault.

    Args:
        root_folder: папка на диске для синхронизации
        force: принудительно перекачать все файлы
        progress_callback: функция(current, total, filename) для прогресса

    Returns:
        dict со статистикой
    """
    if YANDEX_DISK_TOKEN == "YOUR_YANDEX_OAUTH_TOKEN":
        raise ValueError("Укажите YANDEX_DISK_TOKEN в настройках скрипта")

    vault = Path(VAULT_PATH)
    vault.mkdir(parents=True, exist_ok=True)

    # Загружаем состояние
    state = load_sync_state() if not force else {"files": {}, "last_sync": None, "total_synced": 0}

    stats = {
        "new":      0,
        "updated":  0,
        "skipped":  0,
        "errors":   0,
        "total":    0,
        "start_time": datetime.now().isoformat(),
    }

    # Подключаемся к Диску
    log.info("🔌 Подключаюсь к Яндекс Диску...")
    api = YandexDiskAPI(YANDEX_DISK_TOKEN)

    # Получаем список всех файлов
    log.info(f"📂 Сканирую диск: {root_folder}")
    all_files = api.list_all_files(root_folder)
    stats["total"] = len(all_files)

    log.info(f"\n📥 Начинаю синхронизацию ({len(all_files)} файлов)...")

    for i, item in enumerate(all_files, 1):
        filename  = item["name"]
        disk_path = item["path"]

        if progress_callback:
            progress_callback(i, len(all_files), filename)

        # Проверяем нужно ли обновлять
        is_new = disk_path not in state["files"]
        if not force and not file_changed(item, state):
            stats["skipped"] += 1
            log.debug(f"  ⏭  Пропущен: {filename}")
            continue

        try:
            # Скачиваем файл
            log.info(f"  [{i}/{len(all_files)}] ⬇️  {filename} "
                     f"({item.get('size',0)//1024} КБ)")
            data = api.download_file(disk_path)

            # Извлекаем текст
            text = extract_text(filename, data)
            if not text.strip():
                text = "[Файл пуст или не содержит текста]"

            # Сохраняем в vault
            out_path = save_to_vault(vault, item, text)
            log.info(f"  ✅ Сохранён: {out_path.relative_to(vault)}")

            # Обновляем состояние
            state["files"][disk_path] = {
                "md5":      item.get("md5", ""),
                "modified": item.get("modified", ""),
                "synced":   datetime.now().isoformat(),
                "vault_path": str(out_path),
            }

            if is_new:
                stats["new"] += 1
            else:
                stats["updated"] += 1

        except Exception as e:
            log.error(f"  ❌ Ошибка {filename}: {e}")
            stats["errors"] += 1
            continue

        # Сохраняем состояние после каждых 10 файлов
        if i % 10 == 0:
            state["last_sync"] = datetime.now().isoformat()
            state["total_synced"] = state.get("total_synced", 0) + 1
            save_sync_state(state)

    # Создаём индекс раздела
    _create_disk_index(vault, state)

    # Финальное сохранение состояния
    state["last_sync"] = datetime.now().isoformat()
    state["total_synced"] = len(state["files"])
    save_sync_state(state)

    stats["end_time"] = datetime.now().isoformat()
    log.info(f"\n{'='*50}")
    log.info(f"✅ Синхронизация завершена:")
    log.info(f"   Новых:      {stats['new']}")
    log.info(f"   Обновлено:  {stats['updated']}")
    log.info(f"   Пропущено:  {stats['skipped']}")
    log.info(f"   Ошибок:     {stats['errors']}")
    log.info(f"   Всего в базе: {len(state['files'])}")

    return stats


def _create_disk_index(vault: Path, state: dict):
    """Создаёт индексную страницу раздела Яндекс Диска."""
    disk_vault = vault / VAULT_SECTION
    if not disk_vault.exists():
        return

    all_notes = list(disk_vault.rglob("*.md"))
    all_notes = [f for f in all_notes if f.name != "_index.md"]

    # Группируем по папкам
    by_folder = {}
    for note in sorted(all_notes):
        rel    = note.relative_to(disk_vault)
        folder = str(rel.parent) if str(rel.parent) != "." else "/"
        by_folder.setdefault(folder, []).append(note)

    now   = datetime.now().strftime("%d.%m.%Y %H:%M")
    lines = [
        "---",
        "type: disk_index",
        f"updated: \"{now}\"",
        "---",
        "",
        "# 💾 Яндекс Диск",
        "",
        f"Файлов в vault: **{len(all_notes)}**",
        f"Обновлено: {now}",
        f"Всего на диске: {len(state.get('files', {}))}",
        "",
    ]

    for folder, notes in sorted(by_folder.items()):
        lines.append(f"## 📂 {folder}")
        for note in notes[:50]:
            rel = note.relative_to(vault)
            lines.append(f"- [[{rel.with_suffix('')}|{note.stem}]]")
        if len(notes) > 50:
            lines.append(f"_...и ещё {len(notes)-50} файлов_")
        lines.append("")

    (disk_vault / "_index.md").write_text("\n".join(lines), encoding="utf-8")


def get_sync_status() -> dict:
    """Возвращает статус последней синхронизации."""
    state = load_sync_state()
    return {
        "total_files": len(state.get("files", {})),
        "last_sync":   state.get("last_sync"),
        "total_synced": state.get("total_synced", 0),
    }


# ──────────────────────────────────────────────
#  CLI ЗАПУСК
# ──────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Яндекс Диск → Obsidian Vault")
    parser.add_argument("--folder",  default=DISK_ROOT_FOLDER,
                        help="Папка на диске (по умолчанию: /)")
    parser.add_argument("--force",   action="store_true",
                        help="Принудительно перекачать все файлы")
    parser.add_argument("--status",  action="store_true",
                        help="Показать статус последней синхронизации")
    args = parser.parse_args()

    if args.status:
        status = get_sync_status()
        print(f"\n📊 Статус синхронизации Яндекс Диска:")
        print(f"   Файлов в базе: {status['total_files']}")
        print(f"   Последняя синхронизация: {status['last_sync'] or 'никогда'}")
    else:
        sync_disk(root_folder=args.folder, force=args.force)
